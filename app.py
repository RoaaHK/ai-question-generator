from flask import Flask, render_template, request, redirect, url_for, flash, send_file, session, jsonify
import os
import json
import logging
import random
from datetime import datetime
from werkzeug.utils import secure_filename

from pdf_processor import PDFProcessor
from text_splitter import TextSplitter
from db_manager import DBManager
from question_generator import QuestionGenerator

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('app.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

app = Flask(__name__)
app.secret_key = 'your_secret_key'
app.config['SESSION_TYPE'] = 'filesystem'

UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'pdf'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024

OLLAMA_API = os.environ.get('OLLAMA_API', 'http://localhost:11434/api/generate')
OLLAMA_MODEL = os.environ.get('OLLAMA_MODEL', 'llama3.1')

MONGO_URI = os.environ.get('MONGO_URI', 'mongodb://localhost:27017/')
MONGO_DB_NAME = os.environ.get('MONGO_DB_NAME', 'GradProj')

db_manager = DBManager(mongo_uri=MONGO_URI, db_name=MONGO_DB_NAME)
text_splitter = TextSplitter(
    model_name=OLLAMA_MODEL,
    max_allowed_tokens=2048,
    db_connection=db_manager.db
)

question_generator = QuestionGenerator(OLLAMA_API, OLLAMA_MODEL)
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


def process_pdf_complete_pipeline(file_path, file_name, config, page_numbers=None):

    try:
        page_range_hash, is_existing = db_manager.register_file(
            file_path,
            file_name,
            page_numbers=page_numbers
        )
        logger.info(f"File registered: {file_name}, page_range_hash: {page_range_hash[:8]}..., existing: {is_existing}")

        chunks_data = []
        if is_existing:
            existing_chunks = db_manager.get_chunks_by_file_hash(page_range_hash)
            if existing_chunks:
                logger.info(f"Reusing {len(existing_chunks)} existing chunks for this page range")
                chunks_data = existing_chunks

        if not chunks_data:
            logger.info("Processing PDF and creating new chunks...")

            pdf_processor = PDFProcessor()
            extracted_text = pdf_processor.extract_and_preprocess_pdf(
                file_path,
                page_numbers=page_numbers,
                save_to_file=True
            )

            if not extracted_text:
                return None, None, None, False, "Failed to extract text from PDF"

            chunks_data = text_splitter.process_text_direct(
                text=extracted_text,
                file_name=file_name,
                file_hash=page_range_hash,
                min_tokens=150,
                strategy="hierarchical"
            )

            if not chunks_data:
                return None, None, None, False, "Failed to create text chunks"

            success = text_splitter._sync_to_mongodb(chunks_data, collection_name="chunks")
            if not success:
                return None, None, None, False, "Failed to store chunks in database"

            logger.info(f"Created and stored {len(chunks_data)} new chunks")

        session_id = db_manager.create_question_session(
            file_hash=page_range_hash,
            file_name=file_name,
            config=config
        )

        existing_questions = list(db_manager.db.questions.find({
            "session_id": session_id,
            "status": {"$ne": "archived"}
        }))

        if existing_questions:
            logger.info(f"Found {len(existing_questions)} existing questions for session {session_id}")
            return session_id, chunks_data, page_range_hash, True, f"Successfully loaded existing questions for {file_name}"

        logger.info("Generating new questions...")

        if not question_generator.test_connection():
            return None, None, None, False, "Question generation service is not available"

        all_questions = []
        question_types = config["question_types"]
        questions_per_chunk = config["questions_per_chunk"]
        difficulty = config.get("difficulty", "mixed")
        custom_instructions = config.get("custom_instructions", "")

        questions_per_type = {}
        if len(question_types) > 0:
            base_per_type = questions_per_chunk // len(question_types)
            remainder = questions_per_chunk % len(question_types)

            for i, q_type in enumerate(question_types):
                questions_per_type[q_type] = base_per_type + (1 if i < remainder else 0)

        for chunk in chunks_data:
            chunk_id = chunk["chunk_id"]
            chunk_text = chunk["content"]

            logger.info(f"Generating questions for chunk {chunk_id}")

            for q_type, num_questions in questions_per_type.items():
                if num_questions == 0:
                    continue

                chunk_questions = []

                try:
                    if q_type == 'mcq':
                        logger.info(f"Generating {num_questions} MCQ questions for chunk {chunk_id}")
                        mcq_questions = question_generator.generate_mcq(
                            chunk_text,
                            num_questions,
                            difficulty=difficulty,
                            custom_instructions=custom_instructions
                        )
                        for q in mcq_questions:
                            q.update({
                                "question_type": "mcq",
                                "file_name": file_name,
                                "chunk_id": chunk_id,
                                "created_at": datetime.now(),
                                "difficulty": difficulty if difficulty != "mixed" else q.get("difficulty", "medium"),
                                "page_range": config["page_range"]
                            })
                        chunk_questions.extend(mcq_questions)

                    elif q_type == 'true_false':
                        logger.info(f"Generating {num_questions} True/False questions for chunk {chunk_id}")
                        tf_questions = question_generator.generate_true_false(
                            chunk_text,
                            num_questions,
                            difficulty=difficulty,
                            custom_instructions=custom_instructions
                        )
                        for q in tf_questions:
                            q.update({
                                "question_type": "true_false",
                                "file_name": file_name,
                                "chunk_id": chunk_id,
                                "created_at": datetime.now(),
                                "difficulty": difficulty if difficulty != "mixed" else q.get("difficulty", "medium"),
                                "page_range": config["page_range"]
                            })
                        chunk_questions.extend(tf_questions)

                    elif q_type == 'short_answer':
                        logger.info(f"Generating {num_questions} Short Answer questions for chunk {chunk_id}")
                        sa_questions = question_generator.generate_short_answer(
                            chunk_text,
                            num_questions,
                            difficulty=difficulty,
                            custom_instructions=custom_instructions
                        )
                        for q in sa_questions:
                            q.update({
                                "question_type": "short_answer",
                                "file_name": file_name,
                                "chunk_id": chunk_id,
                                "created_at": datetime.now(),
                                "difficulty": difficulty if difficulty != "mixed" else q.get("difficulty", "medium"),
                                "page_range": config["page_range"]
                            })
                        chunk_questions.extend(sa_questions)

                    all_questions.extend(chunk_questions)

                except Exception as e:
                    logger.error(f"Error generating {q_type} questions for chunk {chunk_id}: {e}")

        if all_questions:
            success = db_manager.store_questions_with_metadata(all_questions, session_id)
            if not success:
                return None, None, None, False, "Failed to store generated questions"

            logger.info(f"Generated and stored {len(all_questions)} questions")
        else:
            return None, None, None, False, "No questions could be generated. Please try with different settings."

        logger.info(f"Pipeline completed successfully. Session: {session_id}")
        return session_id, chunks_data, page_range_hash, True, f"Successfully processed {file_name} and generated {len(all_questions)} questions"

    except Exception as e:
        logger.error(f"Pipeline failed: {e}", exc_info=True)
        return None, None, None, False, f"Pipeline error: {str(e)}"


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


@app.route('/')
def index():
    return render_template('index.html')


@app.route('/upload')
def upload_page():
    return render_template('upload.html')


@app.route('/process', methods=['POST'])
def process_upload():
    if 'pdf' not in request.files:
        flash('No file part')
        return redirect(url_for('upload_page'))

    file = request.files['pdf']

    if file.filename == '':
        flash('No file selected')
        return redirect(url_for('upload_page'))

    if file and allowed_file(file.filename):
        try:
            question_types = request.form.getlist('question_types')
            questions_per_chunk = request.form.get('questions_per_chunk', type=int)
            difficulty = request.form.get('difficulty', 'mixed')
            custom_instructions = request.form.get('custom_instructions', '')
            all_pages = 'all_pages' in request.form
            from_page = request.form.get('from_page', type=int) if not all_pages else None
            to_page = request.form.get('to_page', type=int) if not all_pages else None

            if not question_types:
                flash('Please select at least one question type')
                return redirect(url_for('upload_page'))

            if not questions_per_chunk or questions_per_chunk < 1:
                flash('Please enter a valid number of questions (minimum 1)')
                return redirect(url_for('upload_page'))

            if not all_pages and (not from_page or not to_page or from_page > to_page):
                flash('Please enter valid page range')
                return redirect(url_for('upload_page'))

            filename = secure_filename(file.filename)
            base_name = os.path.splitext(filename)[0]
            timestamp = datetime.now().strftime('%Y%m%d%H%M%S')

            if not all_pages and from_page and to_page:
                safe_filename = f"{timestamp}_{base_name}_pages_{from_page}-{to_page}.pdf"
            else:
                safe_filename = f"{timestamp}_{base_name}_all_pages.pdf"

            file_path = os.path.join(app.config['UPLOAD_FOLDER'], safe_filename)
            file.save(file_path)

            page_numbers = None
            if not all_pages and from_page and to_page:
                page_numbers = list(range(from_page, to_page + 1))

            session_config = {
                "question_types": question_types,
                "questions_per_chunk": questions_per_chunk,
                "difficulty": difficulty,
                "custom_instructions": custom_instructions,
                "page_range": {"from": from_page, "to": to_page} if not all_pages else "all",
                "page_count": len(page_numbers) if page_numbers else "all",
                "original_filename": filename
            }

            session_id, chunks_data, page_range_hash, success, message = process_pdf_complete_pipeline(
                file_path=file_path,
                file_name=safe_filename,
                config=session_config,
                page_numbers=page_numbers
            )

            if not success:
                flash(message)
                return redirect(url_for('upload_page'))

            total_questions = 0
            if session_id:
                total_questions = db_manager.db.questions.count_documents({
                    "session_id": session_id,
                    "status": {"$ne": "archived"}
                })

            session['current_session'] = {
                'session_id': session_id,
                'file_hash': page_range_hash,
                'filename': safe_filename,
                'original_name': filename,
                'chunk_count': len(chunks_data) if chunks_data else 0,
                'question_count': total_questions,
                'config': session_config
            }

            return redirect(url_for('questions_by_session', session_id=session_id))

        except Exception as e:
            logger.error(f"Error during upload: {str(e)}")
            flash(f'Error: {str(e)}')
            return redirect(url_for('upload_page'))

    flash('File type not allowed. Please upload a PDF.')
    return redirect(url_for('upload_page'))


@app.route('/questions/session/<session_id>')
def questions_by_session(session_id):
    try:
        session_info = db_manager.db.sessions.find_one({"session_id": session_id})
        if not session_info:
            flash('Session not found.')
            return redirect(url_for('upload_page'))

        default_max = session_info['config'].get('questions_per_chunk', 35)
        max_questions = request.args.get('max', default_max, type=int)

        max_questions = min(max_questions, 35)

        difficulty_filter = request.args.get('difficulty', '')

        filters = {}
        if difficulty_filter:
            filters['difficulty'] = difficulty_filter

        all_questions = db_manager.get_questions_smart(
            session_id,
            max_questions=max_questions,
            filters=filters
        )

        if not all_questions and difficulty_filter:
            total_questions = db_manager.db.questions.count_documents({
                "session_id": session_id,
                "status": {"$ne": "archived"}
            })

            if total_questions > 0:
                return render_template('questions.html',
                                       all_questions=[],
                                       filename=session_info['file_name'],
                                       question_types=session_info['config']['question_types'],
                                       session_id=session_id,
                                       current_session=session_info,
                                       max_questions=max_questions,
                                       total_available=0,
                                       showing_count=0,
                                       no_difficulty_match=True,
                                       requested_difficulty=difficulty_filter)
            else:
                flash('No questions found for this session.')
                return redirect(url_for('upload_page'))

        elif not all_questions:
            flash('No questions found for this session.')
            return redirect(url_for('upload_page'))

        total_available = db_manager.db.questions.count_documents({
            "session_id": session_id,
            "status": {"$ne": "archived"}
        })

        questions_by_chunk = {}
        for question in all_questions:
            chunk_id = question['chunk_id']
            if chunk_id not in questions_by_chunk:
                chunk = text_splitter.find_chunks({"file_name": session_info['file_name'], "chunk_id": chunk_id},
                                                  "chunks")
                chunk_text = chunk[0]['content'] if chunk else "Chunk content not available."

                questions_by_chunk[chunk_id] = {
                    'chunk_id': chunk_id,
                    'chunk_text': chunk_text,
                    'questions': []
                }

            questions_by_chunk[chunk_id]['questions'].append(question)

        grouped_questions = [questions_by_chunk[cid] for cid in sorted(questions_by_chunk.keys())]

        stats = db_manager.get_question_statistics(session_id)

        return render_template('questions.html',
                               all_questions=grouped_questions,
                               filename=session_info['file_name'],
                               question_types=session_info['config']['question_types'],
                               session_id=session_id,
                               current_session=session_info,
                               max_questions=max_questions,
                               total_available=total_available,
                               showing_count=len(all_questions),
                               statistics=stats,
                               default_max=default_max)

    except Exception as e:
        logger.error(f"Error displaying questions: {str(e)}")
        flash(f'Error displaying questions: {str(e)}')
        return redirect(url_for('upload_page'))


@app.route('/questions')
def questions():
    if 'current_session' in session:
        return redirect(url_for('questions_by_session', session_id=session['current_session']['session_id']))
    else:
        return redirect(url_for('all_sessions'))


@app.route('/sessions')
def all_sessions():
    try:
        sessions = list(db_manager.db.sessions.find().sort("created_at", -1))

        for sess in sessions:
            sess['stats'] = db_manager.get_question_statistics(sess['session_id'])

        return render_template('sessions.html', sessions=sessions)

    except Exception as e:
        logger.error(f"Error listing sessions: {str(e)}")
        flash('Error loading sessions.')
        return redirect(url_for('index'))


@app.route('/regenerate/<session_id>/<chunk_id>', methods=['POST'])
def regenerate_chunk_questions(session_id, chunk_id):
    try:
        session_info = db_manager.db.sessions.find_one({"session_id": session_id})
        if not session_info:
            flash('Session not found.')
            return redirect(url_for('questions'))

        chunks = text_splitter.find_chunks({"file_name": session_info['file_name'], "chunk_id": chunk_id}, "chunks")
        if not chunks:
            flash(f'Chunk {chunk_id} not found.')
            return redirect(url_for('questions_by_session', session_id=session_id))

        chunk = chunks[0]

        config = session_info['config']
        question_types = config['question_types']
        questions_per_chunk = config['questions_per_chunk']
        difficulty = config.get('difficulty', 'mixed')
        custom_instructions = config.get('custom_instructions', '')

        if not question_generator.test_connection():
            flash('Question generation service is not available. Please try again later.')
            return redirect(url_for('questions_by_session', session_id=session_id))

        questions_per_type = {}
        if len(question_types) > 0:
            base_per_type = questions_per_chunk // len(question_types)
            remainder = questions_per_chunk % len(question_types)

            for i, q_type in enumerate(question_types):
                questions_per_type[q_type] = base_per_type + (1 if i < remainder else 0)

        new_questions = []
        chunk_text = chunk['content']

        for q_type, num_questions in questions_per_type.items():
            if num_questions == 0:
                continue

            if q_type == 'mcq':
                logger.info(f"Regenerating {num_questions} MCQ questions for chunk {chunk_id}")
                mcq_questions = question_generator.generate_mcq(
                    chunk_text,
                    num_questions,
                    difficulty=difficulty,
                    custom_instructions=custom_instructions
                )
                for q in mcq_questions:
                    q.update({
                        "question_type": "mcq",
                        "file_name": session_info['file_name'],
                        "chunk_id": chunk_id,
                        "created_at": datetime.now(),
                        "difficulty": difficulty if difficulty != "mixed" else None
                    })
                new_questions.extend(mcq_questions)

            elif q_type == 'true_false':
                logger.info(f"Regenerating {num_questions} True/False questions for chunk {chunk_id}")
                tf_questions = question_generator.generate_true_false(
                    chunk_text,
                    num_questions,
                    difficulty=difficulty,
                    custom_instructions=custom_instructions
                )
                for q in tf_questions:
                    q.update({
                        "question_type": "true_false",
                        "file_name": session_info['file_name'],
                        "chunk_id": chunk_id,
                        "created_at": datetime.now(),
                        "difficulty": difficulty if difficulty != "mixed" else None
                    })
                new_questions.extend(tf_questions)

            elif q_type == 'short_answer':
                logger.info(f"Regenerating {num_questions} Short Answer questions for chunk {chunk_id}")
                sa_questions = question_generator.generate_short_answer(
                    chunk_text,
                    num_questions,
                    difficulty=difficulty,
                    custom_instructions=custom_instructions
                )
                for q in sa_questions:
                    q.update({
                        "question_type": "short_answer",
                        "file_name": session_info['file_name'],
                        "chunk_id": chunk_id,
                        "created_at": datetime.now(),
                        "difficulty": difficulty if difficulty != "mixed" else None
                    })
                new_questions.extend(sa_questions)

        keep_best = request.form.get('keep_best', 'true').lower() == 'true'

        if new_questions:
            db_manager.regenerate_questions_smart(session_id, chunk_id, new_questions, keep_best=keep_best)
            logger.info(f"Regenerated {len(new_questions)} questions for chunk {chunk_id}")

            flash(f"Successfully regenerated {len(new_questions)} questions for this chunk.")
        else:
            flash("Failed to generate new questions.")

        return redirect(url_for('questions_by_session', session_id=session_id))

    except Exception as e:
        logger.error(f"Error regenerating questions: {str(e)}")
        flash(f'Error regenerating questions: {str(e)}')
        return redirect(url_for('questions_by_session', session_id=session_id))


@app.route('/download/session/<session_id>/<filetype>')
def download_session(session_id, filetype):
    try:
        session_info = db_manager.db.sessions.find_one({"session_id": session_id})
        if not session_info:
            flash('Session not found.')
            return redirect(url_for('index'))

        all_questions = list(db_manager.db.questions.find({
            "session_id": session_id,
            "status": {"$ne": "archived"}
        }).sort([("chunk_id", 1), ("question_type", 1), ("quality_score", -1)]))

        if not all_questions:
            flash('No questions found for this session.')
            return redirect(url_for('questions_by_session', session_id=session_id))

        questions_by_chunk = {}
        for question in all_questions:
            chunk_id = question['chunk_id']
            if chunk_id not in questions_by_chunk:
                chunks = text_splitter.find_chunks({"file_name": session_info['file_name'], "chunk_id": chunk_id},
                                                   "chunks")
                chunk_text = chunks[0]['content'] if chunks else "Chunk content not available."

                questions_by_chunk[chunk_id] = {
                    'chunk_id': chunk_id,
                    'chunk_text': chunk_text,
                    'questions': []
                }

            questions_by_chunk[chunk_id]['questions'].append(question)

        chunks = [questions_by_chunk[cid] for cid in sorted(questions_by_chunk.keys())]

        if filetype == 'txt':
            output_path = os.path.join(app.config['UPLOAD_FOLDER'], f"session_{session_id}_questions.txt")
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(f"Question Set Generated: {session_info['created_at'].strftime('%Y-%m-%d %H:%M')}\n")
                f.write(f"Configuration: {json.dumps(session_info['config'], indent=2)}\n")
                f.write("=" * 80 + "\n\n")

                for chunk in chunks:
                    f.write(f"CHUNK {chunk['chunk_id']}\n")
                    f.write("=" * 50 + "\n\n")
                    f.write(f"Source Text:\n{chunk['chunk_text']}\n\n")
                    f.write("Questions:\n")

                    for i, q in enumerate(chunk['questions'], 1):
                        q_type = q.get('question_type', 'unknown')
                        difficulty = q.get('difficulty', 'unknown')
                        quality_score = q.get('quality_score', 0)

                        f.write(f"{i}. [{q_type}] [Difficulty: {difficulty}] [Quality: {quality_score}]\n")
                        f.write(f"   {q.get('question', '')}\n")

                        if 'options' in q:
                            for j, option in enumerate(q['options']):
                                opt_letter = chr(65 + j)
                                f.write(f"   {opt_letter}) {option}\n")

                        if 'answer' in q:
                            f.write(f"   Answer: {q['answer']}\n")

                        if 'explanation' in q:
                            f.write(f"   Explanation: {q['explanation']}\n")

                        f.write("\n")

                    f.write("\n" + "=" * 50 + "\n\n")

            return send_file(output_path, as_attachment=True)

        elif filetype == 'docx':
            try:
                from docx import Document
                from docx.shared import Pt, Inches

                doc = Document()

                doc.add_heading(f'Questions Generated from {session_info["file_name"]}', 0)
                doc.add_paragraph(f'Generated: {session_info["created_at"].strftime("%Y-%m-%d %H:%M")}')
                doc.add_paragraph(f'Configuration: {json.dumps(session_info["config"], indent=2)}')

                for chunk in chunks:
                    doc.add_heading(f'Question Set {chunk["chunk_id"]}', 1)

                    p = doc.add_paragraph('Source Text:')
                    p.add_run(chunk['chunk_text']).font.size = Pt(10)

                    doc.add_heading('Questions:', 2)

                    for i, q in enumerate(chunk['questions'], 1):
                        q_type = q.get('question_type', 'unknown')
                        difficulty = q.get('difficulty', 'unknown')
                        quality_score = q.get('quality_score', 0)

                        p = doc.add_paragraph()
                        p.add_run(
                            f"{i}. [{q_type}] [Difficulty: {difficulty}] [Quality: {quality_score}]\n").bold = True
                        p.add_run(q.get('question', ''))

                        if 'options' in q:
                            for j, option in enumerate(q['options']):
                                opt_letter = chr(65 + j)
                                doc.add_paragraph(f"{opt_letter}) {option}", style='List Bullet')

                        if 'answer' in q:
                            p = doc.add_paragraph()
                            p.add_run(f"Answer: {q['answer']}").italic = True

                        if 'explanation' in q:
                            p = doc.add_paragraph()
                            p.add_run(f"Explanation: {q['explanation']}").font.size = Pt(10)

                        doc.add_paragraph()

                    doc.add_page_break()

                output_path = os.path.join(app.config['UPLOAD_FOLDER'], f"session_{session_id}_questions.docx")
                doc.save(output_path)

                return send_file(output_path, as_attachment=True)

            except ImportError:
                flash('python-docx library not installed.')
                return redirect(url_for('questions_by_session', session_id=session_id))

        else:
            flash(f'Unsupported file type: {filetype}')
            return redirect(url_for('questions_by_session', session_id=session_id))

    except Exception as e:
        logger.error(f"Error creating download file: {str(e)}")
        flash(f'Error creating download file: {str(e)}')
        return redirect(url_for('questions_by_session', session_id=session_id))


@app.route('/quiz/session/<session_id>')
def quiz_by_session(session_id):
    try:
        session_info = db_manager.db.sessions.find_one({"session_id": session_id})
        if not session_info:
            flash('Session not found.')
            return redirect(url_for('index'))

        selected_types = request.args.getlist('types')
        if not selected_types:
            selected_types = session_info['config']['question_types']

        default_num = session_info['config'].get('questions_per_chunk', 20)
        num_questions = request.args.get('total', type=int)
        if not num_questions or num_questions <= 0:
            num_questions = min(default_num, 35)

        filters = {"question_types": selected_types}
        quiz_questions = db_manager.get_questions_smart(
            session_id,
            max_questions=num_questions,
            filters=filters
        )

        random.shuffle(quiz_questions)

        return render_template('quiz.html',
                               questions=quiz_questions,
                               filename=session_info['file_name'],
                               question_types=selected_types,
                               session_id=session_id,
                               default_num=default_num)

    except Exception as e:
        logger.error(f"Error creating quiz: {str(e)}")
        flash(f'Error creating quiz: {str(e)}')
        return redirect(url_for('questions_by_session', session_id=session_id))


@app.route('/api/question/<question_id>/answer', methods=['POST'])
def record_answer(question_id):
    try:
        data = request.json
        answered_correctly = data.get('correct', False)

        success = db_manager.update_question_performance(question_id, answered_correctly)

        if success:
            return jsonify({'status': 'success'})
        else:
            return jsonify({'status': 'error', 'message': 'Question not found'}), 404

    except Exception as e:
        logger.error(f"Error recording answer: {str(e)}")
        return jsonify({'status': 'error', 'message': str(e)}), 500


@app.route('/file/<file_name>/sessions')
def file_sessions(file_name):
    try:
        base_name = file_name.split('_pages_')[0] if '_pages_' in file_name else file_name.replace('_all_pages', '')

        sessions = list(db_manager.db.sessions.find({
            "file_name": {"$regex": base_name}
        }).sort("created_at", -1))

        for sess in sessions:
            sess['stats'] = db_manager.get_question_statistics(sess['session_id'])

        return render_template('file_sessions.html',
                               file_name=file_name,
                               sessions=sessions)

    except Exception as e:
        logger.error(f"Error listing file sessions: {str(e)}")
        flash('Error loading sessions.')
        return redirect(url_for('index'))


@app.route('/quiz/<file_name>')
def quiz(file_name):
    file_doc = db_manager.db.files.find_one({"file_name": file_name})
    if not file_doc:
        flash('File not found.')
        return redirect(url_for('index'))

    recent_session = db_manager.db.sessions.find_one(
        {"file_hash": file_doc["file_hash"]},
        sort=[("created_at", -1)]
    )

    if recent_session:
        return redirect(url_for('quiz_by_session', session_id=recent_session["session_id"]))
    else:
        flash('No questions found for this file.')
        return redirect(url_for('index'))


@app.route('/download/<file_name>/<filetype>')
def download(file_name, filetype):
    file_doc = db_manager.db.files.find_one({"file_name": file_name})
    if not file_doc:
        flash('File not found.')
        return redirect(url_for('index'))

    recent_session = db_manager.db.sessions.find_one(
        {"file_hash": file_doc["file_hash"]},
        sort=[("created_at", -1)]
    )

    if recent_session:
        return redirect(url_for('download_session',
                                session_id=recent_session["session_id"],
                                filetype=filetype))
    else:
        flash('No questions found for this file.')
        return redirect(url_for('index'))


@app.errorhandler(413)
def request_entity_too_large(error):
    flash('File too large. Please upload a smaller file (max 16MB).')
    return redirect(url_for('upload_page')), 413


@app.errorhandler(500)
def internal_server_error(error):
    flash('An unexpected error occurred. Please try again later.')
    return redirect(url_for('index')), 500


if __name__ == '__main__':
    app.run(debug=True)